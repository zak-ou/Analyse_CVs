import streamlit as st
import database as db
from app_logic.controller import RecruitmentController
import pandas as pd
import plotly.express as px
import os
import datetime
import importlib
import json
import base64
from docx import Document
from io import BytesIO
importlib.reload(db)

from streamlit_option_menu import option_menu

@st.dialog("Ajouter une offre")
def add_offer_dialog():
    st.write("Remplissez les informations ci-dessous pour créer une nouvelle offre.")
    with st.form("new_job_form"):
        col1, col2 = st.columns(2)
        with col1:
            title = st.text_input("Titre du poste", key="new_job_title")
        with col2:
            domaine = st.text_input("Domaine (ex: IT, Finance...)", key="new_job_domain")
        
        desc = st.text_area("Description du poste", key="new_job_desc")
        skills = st.text_area("Compétences requises (séparées par des virgules)", key="new_job_skills_unique_v2")
        # Ensure unique keys for all inputs

        c1, c2 = st.columns(2)
        with c1:
            exp = st.number_input("Expérience minimum (années)", min_value=0, key="new_job_exp")
        with c2:
            nb_postes = st.number_input("Nombre de postes à pourvoir", min_value=1, value=1, key="new_job_nb")
        
        col3, col4 = st.columns(2)
        with col3:
            deadline_date = st.date_input("Date limite de candidature")
        with col4:
            # Default to current time for standard workday deadlines or immediate
            default_time = datetime.datetime.now().time()
            deadline_time = st.time_input("Heure limite de candidature", value=default_time)
        
        submit = st.form_submit_button("Publier l'offre")
        
        if submit:
            if title and skills and domaine:
                # Combine date and time
                deadline_dt = datetime.datetime.combine(deadline_date, deadline_time)
                deadline = deadline_dt.strftime("%Y-%m-%d %H:%M:%S")
                
                db.create_offer(title, desc, skills, exp, deadline, st.session_state['user_id'], domaine, nb_postes)
                st.success("Offre publiée avec succès !")
                st.rerun()
            else:
                st.error("Le titre, le domaine et les compétences sont obligatoires.")

@st.dialog("Détails de l'opportunité")
def show_market_job_dialog(offer):
    st.markdown(f"""
    <h2 style='color:#6C5CE7; margin-bottom:5px;'>{offer['titre']}</h2>
    <p style='color:#636e72; font-size: 1.1em; margin-top:0;'>📍 {offer['recruteur_prenom']} {offer['recruteur_nom']}</p>
    <div style='display:flex; gap:15px; color:#666; font-size:14px; margin-bottom:20px;'>
        <span>📅 Date limite : <b>{offer['date_limite']}</b></span>
        <span>🎓 Expérience min : <b>{offer['experience_min']} ans</b></span>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("#### 📝 Description du poste")
    st.info(offer['description'])

    st.markdown("#### 🛠 Compétences requises")
    skills_html = "".join([f"<span class='skill-pill'>{s.strip()}</span>" for s in offer['competences_requises'].split(',')])
    st.markdown(f"""
    <style>
    .skill-pill {{
        display: inline-block;
        background-color: #F3F0FF;
        color: #6C5CE7;
        padding: 4px 10px;
        border-radius: 15px;
        font-size: 11px;
        font-weight: 600;
        margin-right: 5px;
        margin-bottom: 5px;
    }}
    </style>
    {skills_html}
    """, unsafe_allow_html=True)

@st.dialog("Visualiseur de CV", width="large")
def view_cv_dialog(cv_path, candidate_name):
    if not cv_path or not os.path.exists(cv_path):
        st.error("Le fichier du CV est introuvable.")
        return
    
    st.subheader(f"CV de {candidate_name}")
    
    file_extension = os.path.splitext(cv_path)[1].lower()
    
    if file_extension == ".pdf":
        try:
            with open(cv_path, "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Erreur lors de l'affichage du PDF : {e}")
            with open(cv_path, "rb") as f:
                st.download_button("Télécharger le CV pour le voir", f, file_name=os.path.basename(cv_path))
    elif file_extension in [".jpg", ".jpeg", ".png"]:
        st.image(cv_path, use_container_width=True)
    else:
        st.info(f"Le format {file_extension} ne peut pas être prévisualisé directement.")
        with open(cv_path, "rb") as f:
            st.download_button("Télécharger le CV", f, file_name=os.path.basename(cv_path))

def render_recruiter_space():
    # --- INTERNAL CSS (Mimicking Candidate Space) ---
    st.markdown("""
    <style>
        .job-title {
            color: #2D3436;
            font-size: 18px;
            font-weight: 700;
            margin-bottom: 5px;
        }
        .job-meta {
            color: #636E72;
            font-size: 13px;
            font-weight: 500;
            display: flex;
            align-items: center;
            gap: 15px;
            margin-bottom: 15px;
        }
        .job-desc {
            font-size: 14px;
            color: #555;
            line-height: 1.5;
            margin-bottom: 15px;
            height: 60px;
            overflow: hidden;
            display: -webkit-box;
            -webkit-line-clamp: 3;
            -webkit-box-orient: vertical;
        }
    </style>
    """, unsafe_allow_html=True)

    # --- SIDEBAR CONFIGURATION ---
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-logo">
            <span class="sidebar-logo-icon">🚀</span>
            <span class="sidebar-logo-text">RecrutIQ</span>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="sidebar-profile">
            <div class="sidebar-avatar">{st.session_state['username'][0].upper()}</div>
            <div class="sidebar-info">
                <div class="sidebar-name">{st.session_state['username']}</div>
                <div class="sidebar-role">{st.session_state['role']}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Primary Navigation (Sidebar)
        sidebar_menu = option_menu(
            menu_title="Menu Principal",
            options=["Tableau de Bord", "Mes Statistiques", "Mon Profil"],
            icons=["house-door-fill", "pie-chart-fill", "person-fill"],
            menu_icon="cast",
            default_index=0,
            styles={
                "container": {"padding": "0!important", "background-color": "transparent"},
                "icon": {"color": "#6C5CE7", "font-size": "1.1rem"}, 
                "nav-link": {"font-size": "15px", "margin":"5px", "border-radius": "8px"},
                "nav-link-selected": {"background-color": "#6C5CE7", "font-weight": "600"},
            }
        )
        
        st.markdown("---")
        if st.button("🚪 Déconnexion", width="stretch", type="secondary"):
            st.session_state['logged_in'] = False
            st.session_state['role'] = None
            st.rerun()

    # --- MAIN CONTENT ROUTING ---
    
    if sidebar_menu == "Tableau de Bord":
        st.header(f"👔 Espace Recruteur")
        
        # Secondary Navigation (Top Horizontal)
        top_menu = option_menu(
            menu_title=None,
            options=["Fil d'actualité", "Mes Offres", "Analyse de Marché"],
            icons=["newspaper", "briefcase-fill", "bar-chart-fill"],
            menu_icon="cast",
            default_index=0,
            orientation="horizontal",
            styles={
                "container": {"padding": "0!important", "background-color": "transparent", "margin-bottom": "20px"},
                "icon": {"color": "#6C5CE7", "font-size": "1.1rem"}, 
                "nav-link": {"font-size": "14px", "margin":"5px", "border-radius": "8px", "padding": "8px"},
                "nav-link-selected": {"background-color": "#6C5CE7", "font-weight": "600"},
            }
        )
        
        # --- TOP MENU CONTENT ---
        
        if top_menu == "Fil d'actualité":
            st.subheader("📰 Fil d'actualité")
            st.caption("Consultez les dernières tendances et offres du marché.")
            
            market_offers = db.get_market_offers()
            
            if not market_offers:
                st.info("Aucune actualité disponible pour le moment.")
            else:
                # --- FILTRES ET RECHERCHE ---
                with st.container(border=True):
                    col_search, col_f1 = st.columns([2, 1])
                    with col_search:
                        search_query = st.text_input("🔍 Rechercher un poste", placeholder="Ex: Data Scientist...")
                    
                    domains = sorted(list(set([o['domaine'] for o in market_offers if o['domaine']])))
                    with col_f1:
                        filter_domain = st.selectbox("Domaine", ["Tous"] + domains)

                st.markdown("---")

                # Filter offers
                filtered_offers = []
                for o in market_offers:
                    if search_query and search_query.lower() not in o['titre'].lower() and search_query.lower() not in o['description'].lower():
                        continue
                    if filter_domain != "Tous" and o['domaine'] != filter_domain:
                        continue
                    filtered_offers.append(o)

                if not filtered_offers:
                    st.info("🔎 Aucune offre ne correspond à vos critères.")
                else:
                    # Grid Layout (2 columns)
                    cols = st.columns(2)
                    for i, offer in enumerate(filtered_offers):
                        with cols[i % 2]:
                            with st.container(border=True):
                                # 1. Titre
                                st.markdown(f"<div class='job-title'>{offer['titre']}</div>", unsafe_allow_html=True)
                                st.caption(f"📍 {offer['recruteur_prenom']} {offer['recruteur_nom']} | {offer['domaine']}")
                                
                                # 2. Métadonnées
                                st.markdown(f"""
                                <div class='job-meta'>
                                    <span>🗓 {offer['date_limite']}</span>
                                    <span>⏳ {offer['experience_min']} ans d'exp.</span>
                                </div>
                                """, unsafe_allow_html=True)
                                
                                # 3. Description
                                st.markdown(f"<div class='job-desc'>{offer['description']}</div>", unsafe_allow_html=True)
                                
                                # 4. Tags
                                help_skills = offer['competences_requises'].split(',')[:2]
                                skills_html = "".join([f"<span class='skill-pill'>{s.strip()}</span>" for s in help_skills])
                                st.markdown(f"<div style='margin-bottom:15px;'>{skills_html}</div>", unsafe_allow_html=True)

                                # 5. Profil Button
                                if st.button("Voir les détails", key=f"market_job_{offer['id']}", width="stretch"):
                                    show_market_job_dialog(offer)

        elif top_menu == "Mes Offres":
            col_header, col_btn = st.columns([0.8, 0.2])
            with col_header:
                st.subheader("📋 Gestion de mes Offres")
            with col_btn:
                if st.button("➕ Ajouter une offre", type="primary", width="stretch"):
                    add_offer_dialog()
            st.markdown("---")
            jobs = db.get_offers(recruteur_id=st.session_state['user_id'])
            
            if not jobs:
                st.info("Vous n'avez publié aucune offre.")
                
            for job in jobs:
                # Add a visual indicator if deadline passed
                deadline_passed = False
                try:
                    d_str = str(job['date_limite'])
                    if len(d_str) > 10:
                        deadline_dt = datetime.datetime.strptime(d_str, "%Y-%m-%d %H:%M:%S")
                        deadline_passed = datetime.datetime.now() > deadline_dt
                    else:
                        deadline_dt = datetime.datetime.strptime(d_str, "%Y-%m-%d")
                        deadline_passed = datetime.date.today() > deadline_dt.date()
                except: pass
                
                status_label = job['statut'].upper()
                if deadline_passed:
                    status_label += " (⏳ DÉLAI DÉPASSÉ)"
                
                with st.expander(f"{job['titre']} ({job['domaine']}) - {status_label}"):
                    # State management for editing
                    is_editing = st.session_state.get(f"edit_{job['id']}", False)

                    if is_editing:
                        st.info("Mode Édition")
                        with st.form(key=f"edit_form_{job['id']}"):
                            col_e1, col_e2 = st.columns(2)
                            with col_e1:
                                u_titre = st.text_input("Titre", value=job['titre'])
                                u_domaine = st.text_input("Domaine", value=job['domaine'])
                            with col_e2:
                                try:
                                    d_str = str(job['date_limite'])
                                    if len(d_str) > 10:
                                        dt_val = datetime.datetime.strptime(d_str, "%Y-%m-%d %H:%M:%S")
                                        d_val = dt_val.date()
                                        t_val = dt_val.time()
                                    else:
                                        d_val = datetime.datetime.strptime(d_str, "%Y-%m-%d").date()
                                        t_val = datetime.time(0, 0)
                                except:
                                    d_val = datetime.date.today()
                                    t_val = datetime.time(0, 0)
                                
                                u_date = st.date_input("Date limite", value=d_val)
                                u_time = st.time_input("Heure limite", value=t_val)
                                u_status = st.selectbox("Statut", ["actif", "clôturé"], index=0 if job['statut'] == 'actif' else 1)
                            
                            u_desc = st.text_area("Description", value=job['description'])
                            u_skills = st.text_area("Compétences", value=job['competences_requises'])
                            c_edit1, c_edit2 = st.columns(2)
                            with c_edit1:
                                u_exp = st.number_input("Expérience Min (années)", value=job['experience_min'] or 0, min_value=0)
                            with c_edit2:
                                current_nb = job['nombre_postes'] if 'nombre_postes' in job.keys() else 1
                                if current_nb is None: current_nb = 1 
                                u_nb_postes = st.number_input("Nombre de postes", value=current_nb, min_value=1)
                            
                            if st.form_submit_button("💾 Enregistrer"):
                                u_deadline_dt = datetime.datetime.combine(u_date, u_time)
                                u_deadline_str = u_deadline_dt.strftime("%Y-%m-%d %H:%M:%S")
                                db.update_offer(job['id'], u_titre, u_desc, u_skills, u_exp, u_deadline_str, u_domaine, u_status, u_nb_postes)
                                st.session_state[f"edit_{job['id']}"] = False
                                st.rerun()
                        
                        if st.button("Annuler", key=f"cancel_{job['id']}"):
                            st.session_state[f"edit_{job['id']}"] = False
                            st.rerun()

                    else:
                        c1, c2 = st.columns([0.85, 0.15])
                        with c1:
                            st.write(f"**Description:** {job['description']}")
                            st.write(f"**Compétences:** {job['competences_requises']}")
                            st.write(f"**Deadline:** {job['date_limite']}")
                        with c2:
                            if st.button("✏️", key=f"btn_edit_{job['id']}", help="Modifier l'offre"):
                                st.session_state[f"edit_{job['id']}"] = True
                                st.rerun()
                    
                    apps = db.get_postulations_for_offer(job['id'])
                    st.write(f"📊 **{len(apps)} Candidature(s)**")
                    
                    if apps:
                        app_data = []
                        for app in apps:
                            skills_detected = "N/A"
                            if app['donnees_analysees']:
                                try:
                                    data = json.loads(app['donnees_analysees'])
                                    skills_detected = ", ".join(data.get('skills', []))
                                except:
                                    skills_detected = "Erreur parsing"

                            app_data.append({
                                "Candidat": f"{app['prenom']} {app['nom']}",
                                "Email": app['email'],
                                "Date": app['date_postulation'],
                                "Score": app['score'] if app['score'] is not None else 0.0,
                                "Compétences": skills_detected,
                                "CV": app['cv_url']
                            })
                        st.dataframe(pd.DataFrame(app_data), width="stretch")

                        if st.button(f"🚀 Analyser CVs", key=f"analyze_{job['id']}"):
                            with st.spinner("Analyse..."):
                                controller = RecruitmentController()
                                required_skills = [s.strip() for s in job['competences_requises'].split(',') if s.strip()]
                                
                                class FileWrapper:
                                    def __init__(self, path):
                                        self.name = os.path.basename(path)
                                        self.path = path
                                    def getbuffer(self):
                                        with open(self.path, "rb") as f:
                                            return f.read()

                                files_to_process = []
                                app_map = {}
                                for app in apps:
                                    f_path = app['cv_url']
                                    if f_path and os.path.exists(f_path):
                                        wrapper = FileWrapper(f_path)
                                        files_to_process.append(wrapper)
                                        app_map[wrapper.name] = app['id']
                                
                                if files_to_process:
                                    results = controller.process_uploads(files_to_process, required_skills)
                                    for res in results:
                                        if res['filename'] in app_map:
                                            db.update_postulation_results(app_map[res['filename']], res['score'], res)
                                    st.success("Analysé!")
                                    st.rerun()
                                else:
                                     st.warning("Aucun fichier trouvé.")

        elif top_menu == "Analyse de Marché":
            st.subheader("📈 Dossier Analyse du Marché du Travail")
            st.caption("Données agrégées basées sur l'ensemble des offres actives de la plateforme.")
            
            market_offers = db.get_market_offers()
            if not market_offers:
                st.info("Aucune donnée disponible pour l'analyse globale.")
            else:
                df_market = pd.DataFrame([dict(row) for row in market_offers])
                
                # --- KPI Row ---
                st.markdown("#### 💎 Indicateurs Clés")
                m1, m2, m3, m4 = st.columns(4)
                m1.metric("Volume d'Offres", len(df_market), delta="+5%" if len(df_market) > 10 else None)
                top_dom = df_market['domaine'].mode()[0] if not df_market.empty else "N/A"
                m2.metric("Secteur Leader", top_dom)
                avg_exp = df_market['experience_min'].mean() if not df_market.empty else 0
                m3.metric("Exp. Moyenne", f"{avg_exp:.1f} ans")
                total_postes = df_market['nombre_postes'].sum() if 'nombre_postes' in df_market.columns else len(df_market)
                m4.metric("Postes Ouverts", int(total_postes))

                st.divider()

                # --- ROW 1: Market Share & Hierarchy ---
                col_left, col_right = st.columns([1, 1])
                
                with col_left:
                    st.markdown("##### 🏛️ Répartition par Secteur (Treemap)")
                    dom_counts = df_market.groupby('domaine').size().reset_index(name='Count')
                    fig_tree = px.treemap(dom_counts, path=['domaine'], values='Count', 
                                         color='Count', color_continuous_scale='Purples',
                                         title="Volume des offres par domaine")
                    fig_tree.update_layout(margin=dict(t=30, l=10, r=10, b=10))
                    st.plotly_chart(fig_tree, width="stretch")
                
                with col_right:
                    st.markdown("##### 🎯 Hiérarchie Postes & Domaines")
                    # Sunburst is great for hierarchical data
                    fig_sun = px.sunburst(df_market, path=['domaine', 'titre'], 
                                         color='domaine', color_discrete_sequence=px.colors.qualitative.Pastel)
                    fig_sun.update_layout(margin=dict(t=30, l=10, r=10, b=10))
                    st.plotly_chart(fig_sun, width="stretch")

                st.divider()

                # --- ROW 2: Experience & Skills ---
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("##### 🎓 Exigences d'Expérience par Domaine")
                    # Box plot shows distribution, median, outliers
                    fig_box = px.box(df_market, x="domaine", y="experience_min", color="domaine",
                                    title="Distribution de l'expérience requise",
                                    points="all", color_discrete_sequence=px.colors.qualitative.Vivid)
                    fig_box.update_layout(showlegend=False, margin=dict(t=40, l=10, r=10, b=10))
                    st.plotly_chart(fig_box, width="stretch")
                
                with col2:
                    st.markdown("##### ⚡ Compétences les plus recherchées")
                    all_skills = []
                    for s in df_market['competences_requises']:
                        if s:
                            all_skills.extend([x.strip() for x in s.split(',')])
                    
                    if all_skills:
                        from collections import Counter
                        top_skills = Counter(all_skills).most_common(12)
                        df_skills = pd.DataFrame(top_skills, columns=['Skill', 'Count'])
                        fig_bar = px.bar(df_skills, x='Count', y='Skill', orientation='h', 
                                        color='Count', color_continuous_scale='Viridis')
                        fig_bar.update_layout(yaxis={'categoryorder':'total ascending'}, 
                                             margin=dict(t=10, l=10, r=10, b=10),
                                             height=400)
                        st.plotly_chart(fig_bar, width="stretch")
                    else:
                        st.info("Données de compétences insuffisantes.")

    # --- SIDEBAR MENU CONTENT (Exclusive Views) ---
    
    elif sidebar_menu == "Mes Statistiques":
        st.header("📊 Mes Statistiques")
        st.markdown("---")
        
        jobs = db.get_offers(recruteur_id=st.session_state['user_id'])
        
        if not jobs:
            st.info("Aucune offre pour afficher les statistiques.")
        else:
            job_titles = {job['titre']: job for job in jobs}
            selected_job_title = st.selectbox("Sélectionnez une offre :", list(job_titles.keys()))
            
            if selected_job_title:
                selected_job = job_titles[selected_job_title]
                apps = db.get_postulations_for_offer(selected_job['id'])
                
                # Statuses are updated by the background worker started in app.py
                # Explicit call removed to improve performance and avoid redundant email processing

                
                # Fetch again to get updated statuses
                apps = db.get_postulations_for_offer(selected_job['id'])
                
                if not apps:
                    st.warning("Aucune candidature.")
                else:
                    app_data = []
                    all_skills = []
                    for app in apps:
                        score = app['score'] if app['score'] is not None else 0
                        skills = []
                        if app['donnees_analysees']:
                            try:
                                data = json.loads(app['donnees_analysees'])
                                if 'skills' in data:
                                     skills = data['skills']
                                     all_skills.extend(skills)
                            except:
                                pass
                        
                        # Logic to determine display status
                        status_db = app['statut_postulation'] 
                        
                        if status_db == "Accepted":
                            etat_display = "✅ Sélectionné"
                        elif status_db == "Refused":
                            etat_display = "❌ Non retenu"
                        else:
                            etat_display = "En attente ⏳"

                        app_data.append({
                            "Candidat": f"{app['prenom']} {app['nom']}",
                            "Email": app['email'],
                            "Date": app['date_postulation'],
                            "Score": score,
                            "Etat": etat_display,
                            "Compétences": ", ".join(skills),
                            "CV": app['cv_url']
                        })
                    
                    df_apps = pd.DataFrame(app_data)
                    df_apps['Date'] = pd.to_datetime(df_apps['Date'])

                    scored_apps = [a for a in app_data if a['Score'] > 0]
                    avg_score = sum(x['Score'] for x in scored_apps) / len(scored_apps) if scored_apps else 0
                    max_score = max(x['Score'] for x in scored_apps) if scored_apps else 0
                    
                    k1, k2, k3 = st.columns(3)
                    k1.metric("Candidats", len(apps))
                    k2.metric("Score Moyen", f"{avg_score:.2f}")
                    k3.metric("Meilleur Score", f"{max_score}")
                    
                    st.markdown("---")
                    st.write("### 📈 Analyse Visuelle")
                    tab1, tab2, tab3 = st.columns(3)
                    with tab1:
                        if not df_apps.empty and scored_apps:
                             fig_hist = px.histogram(df_apps[df_apps['Score'] > 0], x="Score", nbins=20, title="Répartition", color_discrete_sequence=['#6C5CE7'])
                             fig_hist.update_layout(bargap=0.2, showlegend=False, height=300)
                             st.plotly_chart(fig_hist, width="stretch")
                    with tab2:
                        if not df_apps.empty:
                            apps_per_day = df_apps.groupby(df_apps['Date'].dt.date).size().reset_index(name='Count')
                            fig_line = px.line(apps_per_day, x='Date', y='Count', title="Historique", color_discrete_sequence=['#00b894'])
                            fig_line.update_layout(height=300)
                            st.plotly_chart(fig_line, width="stretch")
                    with tab3:
                         if all_skills:
                             from collections import Counter
                             skills_counts = Counter(all_skills).most_common(10)
                             df_skills = pd.DataFrame(skills_counts, columns=['Skill', 'Count'])
                             fig_bar = px.bar(df_skills, x='Count', y='Skill', orientation='h', title="Top Skills", color='Count')
                             fig_bar.update_layout(yaxis={'categoryorder':'total ascending'}, height=300)
                             st.plotly_chart(fig_bar, width="stretch")

                    st.markdown("---")
                    st.write("### 📋 Liste Détaillée des Candidats")
                    
                    # Store app data in session state for the buttons to work reliably if needed
                    st.session_state['current_apps'] = app_data
                    
                    # Display the list with a "Voir CV" action
                    for i, app in enumerate(app_data):
                        with st.container(border=True):
                            c1, c2, c3, c4 = st.columns([0.3, 0.2, 0.3, 0.2])
                            with c1:
                                st.write(f"**{app['Candidat']}**")
                                st.caption(app['Email'])
                            with c2:
                                st.write(f"Score: **{app['Score']:.1f}/100**")
                                st.write(f"Statut: {app['Etat']}")
                            with c3:
                                st.write(f"Compétences: {app['Compétences'][:50]}..." if len(app['Compétences']) > 50 else f"Compétences: {app['Compétences']}")
                            with c4:
                                if st.button("👁️ Voir CV", key=f"view_cv_btn_{i}"):
                                    view_cv_dialog(app['CV'], app['Candidat'])
                    
                    st.markdown("---")
                    # ENHANCED EXPORT
                    # Create a more detailed dataframe for Export
                    export_rows = []
                    for app in apps:
                        details = {}
                        if app['donnees_analysees']:
                            try:
                                details = json.loads(app['donnees_analysees'])
                            except: pass
                        
                        row = {
                            "Candidat": f"{app['prenom']} {app['nom']}",
                            "Email": app['email'],
                            "Date Postulation": app['date_postulation'],
                            "Score Global": app['score'],
                            "Statut Final": app['statut_postulation'],
                            "Compétences Détectées": ", ".join(details.get('skills', [])),
                            "Expérience Détectée": details.get('experience', 'N/A'),
                            "Diplômes": details.get('education', 'N/A'),
                            "Lien CV": app['cv_url']
                        }
                        export_rows.append(row)
                    
                    df_export = pd.DataFrame(export_rows)
                    csv_stat = df_export.to_csv(index=False).encode('utf-8')
                    
                    col_dl1, col_dl2 = st.columns([0.5, 0.5])
                    with col_dl1:
                        st.download_button(
                            label="📥 Télécharger le Rapport Complet (CSV)",
                            data=csv_stat,
                            file_name=f"rapport_complet_{selected_job['titre']}.csv",
                            mime="text/csv",
                            key=f"dl_stat_full_{selected_job['id']}",
                            use_container_width=True
                        )
                    with col_dl2:
                        # WORD REPORT GENERATION
                        doc = Document()
                        doc.add_heading(f"Rapport de Recrutement : {selected_job['titre']}", 0)
                        
                        doc.add_heading("Résumé de l'Offre", level=1)
                        p = doc.add_paragraph()
                        p.add_run(f"Domaine : ").bold = True
                        p.add_run(f"{selected_job['domaine']}\n")
                        p.add_run(f"Statut : ").bold = True
                        p.add_run(f"{selected_job['statut']}\n")
                        p.add_run(f"Nombre de postes : ").bold = True
                        p.add_run(f"{selected_job['nombre_postes']}\n")
                        p.add_run(f"Date du rapport : ").bold = True
                        p.add_run(f"{datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")

                        doc.add_heading("Statistiques Clés", level=1)
                        table_stats = doc.add_table(rows=1, cols=3)
                        table_stats.style = 'Table Grid'
                        hdr_cells = table_stats.rows[0].cells
                        hdr_cells[0].text = 'Total Candidatures'
                        hdr_cells[1].text = 'Score Moyen'
                        hdr_cells[2].text = 'Meilleur Score'
                        
                        row_cells = table_stats.add_row().cells
                        row_cells[0].text = str(len(apps))
                        row_cells[1].text = f"{avg_score:.2f}/100"
                        row_cells[2].text = f"{max_score:.2f}/100"

                        doc.add_heading("Détails des Candidats", level=1)
                        table_cand = doc.add_table(rows=1, cols=4)
                        table_cand.style = 'Table Grid'
                        hdr_c = table_cand.rows[0].cells
                        hdr_c[0].text = 'Nom du Candidat'
                        hdr_c[1].text = 'Score'
                        hdr_c[2].text = 'Statut'
                        hdr_c[3].text = 'Compétences Principales'
                        
                        for row in export_rows:
                            cells = table_cand.add_row().cells
                            cells[0].text = row['Candidat']
                            cells[1].text = f"{row['Score Global'] if row['Score Global'] is not None else 0:.1f}/100"
                            cells[2].text = row['Statut Final']
                            cells[3].text = row['Compétences Détectées'][:100] # Truncate for table
                        
                        # Save to stream
                        doc_io = BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        st.download_button(
                            label="� Télécharger le Rapport (Word)",
                            data=doc_io,
                            file_name=f"rapport_{selected_job['titre']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_stat_docx_{selected_job['id']}",
                            use_container_width=True
                        )

                    
    elif sidebar_menu == "Mon Profil":
        import services.profile_service as ps
        
        # Fresh fetch
        user_data = db.get_user_by_id(st.session_state['user_id'], st.session_state['role'])
        if user_data:
            user_data = dict(user_data)
        completion = ps.calculate_recruiter_completion(user_data)
        
        # 1. PROFILE HEADER
        avatar_html = user_data['prenom'][0].upper()
        if user_data.get('photo_url') and os.path.exists(user_data['photo_url']):
            try:
                with open(user_data["photo_url"], "rb") as img_file:
                    b64_content = base64.b64encode(img_file.read()).decode()
                    avatar_html = f'<img src="data:image/png;base64,{b64_content}" />'
            except:
                pass

        st.markdown(f"""
<div class="profile-header-container">
<div class="profile-avatar-large">
{avatar_html}
</div>
<div class="profile-name">{user_data['prenom']} {user_data['nom']}</div>
<span class="profile-role-badge">Recruteur Vérifié</span>
<div class="profile-progress-container">
<div style="display: flex; justify-content: space-between; margin-bottom: 5px; font-size: 0.85rem;">
<span>Complétion du profil</span>
<span>{completion}%</span>
</div>
<div class="profile-progress-bar">
<div class="profile-progress-fill" style="width: {completion}%"></div>
</div>
</div>
</div>
""", unsafe_allow_html=True)

        # 2. CONTENT TABS
        r_tab1, r_tab2, r_tab3 = st.tabs(["👤 Infos Personnelles", "🏢 Entreprise", "⚙️ Paramètres"])
        
        with r_tab1:
            with st.form("edit_recruiter_main"):
                col1, col2 = st.columns(2)
                with col1:
                    u_prenom = st.text_input("Prénom", value=user_data['prenom'])
                    u_nom = st.text_input("Nom", value=user_data['nom'])
                with col2:
                    u_email = st.text_input("Email", value=user_data['email'], disabled=True)
                    u_tele = st.text_input("Numéro de Téléphone", value=user_data['num_tele'] or "")
                
                u_photo = st.file_uploader("Modifier la photo de profil", type=['png', 'jpg', 'jpeg'])
                
                if st.form_submit_button("Sauvegarder les informations", width="stretch", type="primary"):
                    photo_path = user_data['photo_url']
                    if u_photo:
                        os.makedirs("data/profiles", exist_ok=True)
                        photo_path = f"data/profiles/avatar_rec_{st.session_state['user_id']}.png"
                        with open(photo_path, "wb") as f:
                            f.write(u_photo.getbuffer())
                    
                    db.update_recruiter_profile(
                        st.session_state['user_id'],
                        u_nom,
                        u_prenom,
                        u_tele,
                        photo_path,
                        user_data['entreprise_nom'],
                        user_data['entreprise_site'],
                        user_data['entreprise_description']
                    )
                    st.success("Profil mis à jour !")
                    st.rerun()

        with r_tab2:
            with st.form("edit_recruiter_company"):
                u_ent_nom = st.text_input("Nom de l'entreprise", value=user_data['entreprise_nom'] or "")
                u_ent_site = st.text_input("Site web", value=user_data['entreprise_site'] or "", placeholder="https://www.entreprise.com")
                u_ent_desc = st.text_area("Description de l'entreprise", value=user_data['entreprise_description'] or "", height=150)
                
                if st.form_submit_button("Mettre à jour les infos entreprise", width="stretch", type="primary"):
                    db.update_recruiter_profile(
                        st.session_state['user_id'],
                        user_data['nom'],
                        user_data['prenom'],
                        user_data['num_tele'],
                        user_data['photo_url'],
                        u_ent_nom,
                        u_ent_site,
                        u_ent_desc
                    )
                    st.success("Informations entreprise mises à jour !")
                    st.rerun()

        with r_tab3:
            st.info("Les paramètres de sécurité seront bientôt disponibles.")
            if st.button("🚪 Déconnexion", width="stretch", type="secondary"):
                st.session_state['logged_in'] = False
                st.rerun()
